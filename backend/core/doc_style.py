"""
Shared investor-grade styling for generated PDF and DOCX artifacts
(pitch decks, executive summaries, due diligence Q&A).

Keeps a single consistent look (cover page, colored headings, rules,
header/footer with page numbers) across every document-generating tool.
"""

from datetime import datetime

# ---------------------------------------------------------------------------
# Palette — dark navy + gold accent, the classic "investor deck" combo.
# ---------------------------------------------------------------------------
NAVY_HEX = "132A46"
ACCENT_HEX = "C9973B"
SLATE_HEX = "5B6B7F"
LIGHT_RULE_HEX = "D9DEE5"
BODY_TEXT_HEX = "2A2A2A"


# ---------------------------------------------------------------------------
# ReportLab (PDF) helpers
# ---------------------------------------------------------------------------

def get_pdf_styles():
    """Return a dict of ParagraphStyles sharing one investor-deck look."""
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_JUSTIFY
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

    navy = colors.HexColor(f"#{NAVY_HEX}")
    accent = colors.HexColor(f"#{ACCENT_HEX}")
    slate = colors.HexColor(f"#{SLATE_HEX}")
    body_color = colors.HexColor(f"#{BODY_TEXT_HEX}")

    base = getSampleStyleSheet()

    return {
        "cover_title": ParagraphStyle(
            name="CoverTitle",
            parent=base["Title"],
            fontName="Helvetica-Bold",
            fontSize=30,
            leading=34,
            textColor=navy,
            spaceAfter=10,
            alignment=0,
        ),
        "cover_subtitle": ParagraphStyle(
            name="CoverSubtitle",
            parent=base["Normal"],
            fontName="Helvetica-Oblique",
            fontSize=13,
            leading=17,
            textColor=slate,
            spaceAfter=6,
        ),
        "cover_meta": ParagraphStyle(
            name="CoverMeta",
            parent=base["Normal"],
            fontName="Helvetica",
            fontSize=9.5,
            leading=13,
            textColor=slate,
        ),
        "h1": ParagraphStyle(
            name="SectionHeading",
            parent=base["Heading1"],
            fontName="Helvetica-Bold",
            fontSize=13.5,
            leading=17,
            textColor=navy,
            spaceBefore=18,
            spaceAfter=2,
        ),
        "h2": ParagraphStyle(
            name="SubHeading",
            parent=base["Heading2"],
            fontName="Helvetica-Bold",
            fontSize=11.5,
            leading=15,
            textColor=accent,
            spaceBefore=12,
            spaceAfter=4,
        ),
        "body": ParagraphStyle(
            name="Body",
            parent=base["Normal"],
            fontName="Helvetica",
            fontSize=10.3,
            leading=15,
            textColor=body_color,
            spaceAfter=8,
            alignment=TA_JUSTIFY,
        ),
    }


def make_page_decorator(doc_title: str, footer_label: str = ""):
    """
    Build (on_first_page, on_later_pages) callbacks for SimpleDocTemplate.build().
    First page renders a clean cover accent (no repeated header). Later pages get
    a slim header rule + running title, and every page gets a footer with page number.
    """
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.units import inch

    navy = colors.HexColor(f"#{NAVY_HEX}")
    accent = colors.HexColor(f"#{ACCENT_HEX}")
    slate = colors.HexColor(f"#{SLATE_HEX}")
    page_w, page_h = letter

    def _footer(canvas, doc):
        canvas.saveState()
        canvas.setStrokeColor(colors.HexColor(f"#{LIGHT_RULE_HEX}"))
        canvas.setLineWidth(0.6)
        canvas.line(0.75 * inch, 0.62 * inch, page_w - 0.75 * inch, 0.62 * inch)
        canvas.setFont("Helvetica", 8)
        canvas.setFillColor(slate)
        canvas.drawString(0.75 * inch, 0.45 * inch, footer_label or doc_title)
        canvas.drawRightString(page_w - 0.75 * inch, 0.45 * inch, f"Page {doc.page}")
        canvas.restoreState()

    def on_first_page(canvas, doc):
        canvas.saveState()
        canvas.setFillColor(navy)
        canvas.rect(0, page_h - 0.22 * inch, page_w, 0.22 * inch, fill=1, stroke=0)
        canvas.setFillColor(accent)
        canvas.rect(0, page_h - 0.28 * inch, page_w, 0.06 * inch, fill=1, stroke=0)
        canvas.restoreState()
        _footer(canvas, doc)

    def on_later_pages(canvas, doc):
        canvas.saveState()
        canvas.setStrokeColor(accent)
        canvas.setLineWidth(1.4)
        canvas.line(0.75 * inch, page_h - 0.65 * inch, page_w - 0.75 * inch, page_h - 0.65 * inch)
        canvas.setFont("Helvetica-Bold", 8.5)
        canvas.setFillColor(navy)
        canvas.drawString(0.75 * inch, page_h - 0.55 * inch, doc_title.upper())
        canvas.restoreState()
        _footer(canvas, doc)

    return on_first_page, on_later_pages


def escape_pdf_text(text: str) -> str:
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def cover_flowables(title: str, subtitle: str, meta_lines: list[str] | None = None):
    """Flowables for a clean cover block at the top of the first page."""
    from reportlab.lib.units import inch
    from reportlab.platypus import Paragraph, Spacer

    styles = get_pdf_styles()
    flow = [Spacer(1, 0.55 * inch), Paragraph(escape_pdf_text(title), styles["cover_title"])]
    if subtitle:
        flow.append(Paragraph(escape_pdf_text(subtitle), styles["cover_subtitle"]))
    flow.append(Spacer(1, 0.08 * inch))
    for line in meta_lines or []:
        flow.append(Paragraph(escape_pdf_text(line), styles["cover_meta"]))
    flow.append(Spacer(1, 0.35 * inch))
    return flow


def today_str() -> str:
    return datetime.now().strftime("%B %d, %Y")


# ---------------------------------------------------------------------------
# python-docx helpers
# ---------------------------------------------------------------------------

def style_docx_base(doc):
    """Apply base fonts/colors to a freshly created Document()."""
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(BODY_TEXT_HEX)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.25

    title = doc.styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(30)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string(NAVY_HEX)

    h1 = doc.styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(15)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor.from_string(NAVY_HEX)
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(4)

    h2 = doc.styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(12.5)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor.from_string(ACCENT_HEX)
    h2.paragraph_format.space_before = Pt(10)
    h2.paragraph_format.space_after = Pt(3)

    section = doc.sections[0]
    section.left_margin = Pt(62)
    section.right_margin = Pt(62)
    section.top_margin = Pt(54)
    section.bottom_margin = Pt(54)


def add_bottom_border(paragraph, color_hex: str = ACCENT_HEX, size: int = 12):
    """Add a horizontal rule under a paragraph via a bottom border on its XML."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color_hex)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_heading_with_rule(doc, text: str, level: int = 1, color_hex: str = ACCENT_HEX):
    heading = doc.add_heading(text, level=level)
    add_bottom_border(heading, color_hex=color_hex, size=10 if level == 1 else 6)
    return heading


def add_cover_block(doc, title: str, subtitle: str, meta_lines: list[str] | None = None):
    """Title + subtitle + accent rule + meta lines, styled like a deck cover."""
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    t = doc.add_paragraph()
    run = t.add_run(title)
    run.font.name = "Calibri"
    run.font.size = Pt(30)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(NAVY_HEX)

    if subtitle:
        s = doc.add_paragraph()
        srun = s.add_run(subtitle)
        srun.font.name = "Calibri"
        srun.font.size = Pt(13)
        srun.italic = True
        srun.font.color.rgb = RGBColor.from_string(SLATE_HEX)

    rule_p = doc.add_paragraph()
    rule_p.paragraph_format.space_after = Pt(10)
    add_bottom_border(rule_p, color_hex=ACCENT_HEX, size=16)

    for line in meta_lines or []:
        m = doc.add_paragraph()
        mrun = m.add_run(line)
        mrun.font.name = "Calibri"
        mrun.font.size = Pt(9.5)
        mrun.font.color.rgb = RGBColor.from_string(SLATE_HEX)

    doc.add_paragraph()


def add_page_numbers(doc):
    """Insert a centered 'Page X of Y' field in the footer."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    footer = doc.sections[0].footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.text = ""

    def _field(instr: str):
        run = p.add_run()
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string(SLATE_HEX)
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = instr
        fld_sep = OxmlElement("w:fldChar")
        fld_sep.set(qn("w:fldCharType"), "separate")
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        r_elem = run._r
        r_elem.append(fld_begin)
        r_elem.append(instr_text)
        r_elem.append(fld_sep)
        r_elem.append(fld_end)

    lead = p.add_run("Page ")
    lead.font.size = Pt(9)
    lead.font.color.rgb = RGBColor.from_string(SLATE_HEX)
    _field("PAGE")
    mid = p.add_run(" of ")
    mid.font.size = Pt(9)
    mid.font.color.rgb = RGBColor.from_string(SLATE_HEX)
    _field("NUMPAGES")

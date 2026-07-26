"""
Deck Creator tools — generate pitch deck / product report as PDF or DOCX.
"""

import json
import os
import re
from datetime import datetime

from core.config import config

_ARTIFACTS_DIR = getattr(config, "artifacts_root_dir", "./artifacts")

SECTION_ORDER = [
    "problem",
    "solution",
    "market_size",
    "product",
    "traction",
    "business_model",
    "gtm_strategy",
    "competition",
]

SECTION_TITLES = {
    "problem": "Problem",
    "solution": "Solution",
    "market_size": "Market Size",
    "product": "Product",
    "traction": "Traction",
    "business_model": "Business Model",
    "gtm_strategy": "GTM Strategy",
    "competition": "Competition",
}


def _sanitize_filename(name: str) -> str:
    stem = re.sub(r"[^\w\s-]", "", (name or "deck").strip())[:80]
    return stem or "deck"


_PLACEHOLDER_PHRASES = frozenset(
    s.lower() for s in (
        "to be added", "tbd", "t.b.d.", "n/a", "na", "none", "—", "placeholder",
        "add later", "to be filled", "coming soon", "not yet provided", "not provided",
    )
)


def _parse_content(content_json: str) -> tuple[str, dict]:
    """Parse content_json and return (company_name, sections dict). Skip placeholders."""
    try:
        data = json.loads(content_json)
    except json.JSONDecodeError:
        return "", {}
    company = (data.get("company_name") or data.get("company") or "Product Deck").strip()
    sections = {}
    for key in SECTION_ORDER:
        val = data.get(key)
        if val is None:
            continue
        text = str(val).strip()
        if not text:
            continue
        if text.lower() in _PLACEHOLDER_PHRASES:
            continue
        if len(text) < 4 and text.lower() in ("n/a", "na", "tbd"):
            continue
        sections[key] = text
    return company, sections


def create_deck_pdf(content_json: str, company_name: str = "") -> str:
    """
    Create a pitch deck / product report as a PDF and save it to the artifacts directory.

    Call this when the user wants the report in PDF format. Build the content from conversation
    context (problem, solution, market size, product, traction, business model, GTM, competition).

    Args:
        content_json: JSON string with keys: company_name, problem, solution, market_size,
                     product, traction, business_model, gtm_strategy, competition.
                     Each value is a string (can be multi-line). Omit or empty string to skip.
        company_name: Override for document title/filename (optional if in content_json).

    Returns:
        Message with download filename, e.g. "Deck PDF created. Download: deck_Company_20250228.pdf"
    """
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.units import inch
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, HRFlowable
    except ImportError:
        return "PDF creation failed: reportlab is not installed. Install with: pip install reportlab"

    from core.doc_style import (
        get_pdf_styles, make_page_decorator, escape_pdf_text, cover_flowables,
        today_str, ACCENT_HEX,
    )

    company, sections = _parse_content(content_json)
    if not company and company_name:
        company = company_name.strip() or "Product Deck"
    if not company:
        company = "Product Deck"

    if not sections:
        return "No content provided. Pass a JSON with at least one of: problem, solution, market_size, product, traction, business_model, gtm_strategy, competition."

    os.makedirs(_ARTIFACTS_DIR, exist_ok=True)
    safe_name = _sanitize_filename(company)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"deck_{safe_name}_{timestamp}.pdf"
    filepath = os.path.join(_ARTIFACTS_DIR, filename)

    try:
        doc = SimpleDocTemplate(
            filepath,
            pagesize=letter,
            leftMargin=0.75 * inch,
            rightMargin=0.75 * inch,
            topMargin=0.85 * inch,
            bottomMargin=0.85 * inch,
        )
        styles = get_pdf_styles()

        story = cover_flowables(
            company,
            "Pitch Deck & Product Report",
            [today_str(), "Prepared for investor review"],
        )

        for key in SECTION_ORDER:
            if key not in sections:
                continue
            title = SECTION_TITLES.get(key, key.replace("_", " ").title())
            story.append(Paragraph(escape_pdf_text(title), styles["h1"]))
            story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor(f"#{ACCENT_HEX}"), spaceAfter=8))
            text = escape_pdf_text(sections[key])
            story.append(Paragraph(text.replace("\n", "<br/>"), styles["body"]))
            story.append(Spacer(1, 0.06 * inch))

        on_first, on_later = make_page_decorator(company, footer_label=f"{company} — Confidential")
        doc.build(story, onFirstPage=on_first, onLaterPages=on_later)
        return f"Deck PDF created. Download: {filename}"
    except Exception as e:
        return f"PDF creation failed: {str(e)}"


def create_deck_docx(content_json: str, company_name: str = "") -> str:
    """
    Create a pitch deck / product report as a DOCX and save it to the artifacts directory.

    Call this when the user wants the report in DOCX format. Build the content from conversation
    context (problem, solution, market size, product, traction, business model, GTM, competition).

    Args:
        content_json: JSON string with keys: company_name, problem, solution, market_size,
                     product, traction, business_model, gtm_strategy, competition.
                     Each value is a string (can be multi-line). Omit or empty string to skip.
        company_name: Override for document title/filename (optional if in content_json).

    Returns:
        Message with download filename, e.g. "Deck DOCX created. Download: deck_Company_20250228.docx"
    """
    try:
        from docx import Document
    except ImportError:
        return "DOCX creation failed: python-docx is not installed. Install with: pip install python-docx"

    from core.doc_style import style_docx_base, add_cover_block, add_heading_with_rule, add_page_numbers, today_str

    company, sections = _parse_content(content_json)
    if not company and company_name:
        company = company_name.strip() or "Product Deck"
    if not company:
        company = "Product Deck"

    if not sections:
        return "No content provided. Pass a JSON with at least one of: problem, solution, market_size, product, traction, business_model, gtm_strategy, competition."

    os.makedirs(_ARTIFACTS_DIR, exist_ok=True)
    safe_name = _sanitize_filename(company)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"deck_{safe_name}_{timestamp}.docx"
    filepath = os.path.join(_ARTIFACTS_DIR, filename)

    try:
        doc = Document()
        style_docx_base(doc)
        add_cover_block(
            doc,
            company,
            "Pitch Deck & Product Report",
            [today_str(), "Prepared for investor review"],
        )
        add_page_numbers(doc)

        for key in SECTION_ORDER:
            if key not in sections:
                continue
            title = SECTION_TITLES.get(key, key.replace("_", " ").title())
            add_heading_with_rule(doc, title, level=1)
            for para in sections[key].strip().split("\n\n"):
                if para.strip():
                    doc.add_paragraph(para.strip())

        doc.save(filepath)
        return f"Deck DOCX created. Download: {filename}"
    except Exception as e:
        return f"DOCX creation failed: {str(e)}"
